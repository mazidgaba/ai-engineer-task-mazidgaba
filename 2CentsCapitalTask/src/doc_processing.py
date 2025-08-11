"""
Document Processing Module
Handles .docx parsing, type detection, and compliance checking
"""

import docx
import json
import re
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import logging
from checklist import ADGM_CHECKLISTS, verify_checklist_completeness

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        # Document type detection patterns
        self.document_patterns = {
            "Articles of Association": [
                r"articles\s+of\s+association",
                r"aoa",
                r"company\s+constitution",
                r"internal\s+regulations"
            ],
            "Memorandum of Association": [
                r"memorandum\s+of\s+association",
                r"moa",
                r"company\s+objectives",
                r"business\s+purpose"
            ],
            "UBO Declaration": [
                r"ubo\s+declaration",
                r"ultimate\s+beneficial\s+owner",
                r"beneficial\s+ownership",
                r"ownership\s+structure"
            ],
            "Board Resolution": [
                r"board\s+resolution",
                r"directors\s+resolution",
                r"board\s+decision",
                r"directors\s+meeting"
            ],
            "Employment Contract": [
                r"employment\s+contract",
                r"employment\s+agreement",
                r"staff\s+contract",
                r"employee\s+terms"
            ],
            "Compliance Manual": [
                r"compliance\s+manual",
                r"compliance\s+policy",
                r"regulatory\s+compliance",
                r"compliance\s+framework"
            ]
        }
        
        # ADGM compliance rules
        self.compliance_rules = {
            "jurisdiction": {
                "required": ["ADGM", "Abu Dhabi Global Market", "ADGM Courts"],
                "forbidden": ["UAE Federal Courts", "Dubai Courts", "Dubai International Financial Centre"],
                "severity": "High"
            },
            "governing_law": {
                "required": ["ADGM", "English Law", "Common Law"],
                "forbidden": ["UAE Federal Law", "Dubai Law"],
                "severity": "High"
            },
            "signature_blocks": {
                "required": ["signature", "witness", "notary"],
                "severity": "Medium"
            },
            "document_structure": {
                "required": ["clause", "section", "article"],
                "severity": "Low"
            }
        }

    def parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Parse a .docx file and extract content and metadata"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text content
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            full_text = "\n".join(content)
            
            # Extract metadata
            metadata = {
                "filename": Path(file_path).name,
                "word_count": len(full_text.split()),
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables),
                "file_size": Path(file_path).stat().st_size
            }
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise

    def detect_document_type(self, content: str, filename: str) -> Tuple[str, str, float]:
        """Detect document type with confidence score"""
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        best_match = None
        best_score = 0.0
        best_category = "Unknown"
        
        # Check each document pattern
        for doc_type, patterns in self.document_patterns.items():
            score = 0.0
            
            # Pattern matching in content
            for pattern in patterns:
                if re.search(pattern, content_lower):
                    score += 0.6  # Content match weight
                
                # Filename matching
                if re.search(pattern, filename_lower):
                    score += 0.4  # Filename match weight
            
            # Additional keyword scoring
            if doc_type == "Articles of Association":
                if "clause" in content_lower and "company" in content_lower:
                    score += 0.2
            elif doc_type == "Memorandum of Association":
                if "object" in content_lower and "purpose" in content_lower:
                    score += 0.2
            elif doc_type == "UBO Declaration":
                if "ownership" in content_lower and "percentage" in content_lower:
                    score += 0.2
            
            if score > best_score:
                best_score = score
                best_match = doc_type
        
        # Determine category based on document type
        if best_match in ["Articles of Association", "Memorandum of Association", "UBO Declaration"]:
            best_category = "Company Incorporation"
        elif best_match in ["Employment Contract", "Confidentiality Agreement"]:
            best_category = "Employment"
        elif best_match in ["Compliance Manual", "Risk Management Policy"]:
            best_category = "Compliance"
        else:
            best_category = "General Legal"
        
        return best_match or "Unknown Document", best_category, min(best_score, 1.0)

    def check_adgm_compliance(self, content: str, document_type: str) -> List[Dict]:
        """Check document for ADGM compliance issues"""
        issues = []
        content_lower = content.lower()
        
        # Check jurisdiction compliance
        jurisdiction_issue = self._check_jurisdiction_compliance(content_lower)
        if jurisdiction_issue:
            issues.append(jurisdiction_issue)
        
        # Check governing law
        governing_law_issue = self._check_governing_law_compliance(content_lower)
        if governing_law_issue:
            issues.append(governing_law_issue)
        
        # Check signature blocks
        signature_issue = self._check_signature_compliance(content_lower)
        if signature_issue:
            issues.append(signature_issue)
        
        # Check document structure
        structure_issue = self._check_structure_compliance(content_lower, document_type)
        if structure_issue:
            issues.append(structure_issue)
        
        # Document-specific checks
        doc_specific_issues = self._check_document_specific_compliance(content_lower, document_type)
        issues.extend(doc_specific_issues)
        
        return issues

    def _check_jurisdiction_compliance(self, content: str) -> Optional[Dict]:
        """Check if jurisdiction clauses are ADGM-compliant"""
        rule = self.compliance_rules["jurisdiction"]
        
        # Check for forbidden jurisdictions
        for forbidden in rule["forbidden"]:
            if forbidden.lower() in content:
                return {
                    "issue": f"Incorrect jurisdiction: {forbidden}",
                    "severity": rule["severity"],
                    "suggestion": "Update jurisdiction to ADGM Courts",
                    "regulation": "ADGM Companies Regulations 2020, Art. 6",
                    "section": "Jurisdiction"
                }
        
        # Check for required ADGM jurisdiction
        has_adgm_jurisdiction = any(required.lower() in content for required in rule["required"])
        if not has_adgm_jurisdiction:
            return {
                "issue": "Missing ADGM jurisdiction clause",
                "severity": rule["severity"],
                "suggestion": "Add jurisdiction clause specifying ADGM Courts",
                "regulation": "ADGM Companies Regulations 2020, Art. 6",
                "section": "Jurisdiction"
            }
        
        return None

    def _check_governing_law_compliance(self, content: str) -> Optional[Dict]:
        """Check if governing law is ADGM-compliant"""
        rule = self.compliance_rules["governing_law"]
        
        # Check for forbidden laws
        for forbidden in rule["forbidden"]:
            if forbidden.lower() in content:
                return {
                    "issue": f"Incorrect governing law: {forbidden}",
                    "severity": rule["severity"],
                    "suggestion": "Update governing law to ADGM/English Law",
                    "regulation": "ADGM Companies Regulations 2020, Art. 7",
                    "section": "Governing Law"
                }
        
        return None

    def _check_signature_compliance(self, content: str) -> Optional[Dict]:
        """Check for proper signature blocks"""
        rule = self.compliance_rules["signature_blocks"]
        
        has_signature = any(required.lower() in content for required in rule["required"])
        if not has_signature:
            return {
                "issue": "Missing signature section",
                "severity": rule["severity"],
                "suggestion": "Add proper signature blocks with witness section",
                "regulation": "ADGM Companies Regulations 2020, Art. 12",
                "section": "Execution"
            }
        
        return None

    def _check_structure_compliance(self, content: str, document_type: str) -> Optional[Dict]:
        """Check document structure compliance"""
        rule = self.compliance_rules["document_structure"]
        
        has_structure = any(required.lower() in content for required in rule["required"])
        if not has_structure:
            return {
                "issue": "Poor document structure",
                "severity": rule["severity"],
                "suggestion": "Organize document with proper clauses and sections",
                "regulation": "ADGM Companies Regulations 2020, Art. 10",
                "section": "Document Requirements"
            }
        
        return None

    def _check_document_specific_compliance(self, content: str, document_type: str) -> List[Dict]:
        """Check document-specific compliance requirements"""
        issues = []
        
        if document_type == "Articles of Association":
            if "share capital" not in content and "capital" not in content:
                issues.append({
                    "issue": "Missing share capital information",
                    "severity": "Medium",
                    "suggestion": "Include share capital structure and classes",
                    "regulation": "ADGM Companies Regulations 2020, Art. 15",
                    "section": "Share Capital"
                })
        
        elif document_type == "UBO Declaration":
            if "percentage" not in content and "ownership" not in content:
                issues.append({
                    "issue": "Missing ownership percentages",
                    "severity": "High",
                    "suggestion": "Include beneficial ownership percentages",
                    "regulation": "ADGM Companies Regulations 2020, Art. 20",
                    "section": "Beneficial Ownership"
                })
        
        elif document_type == "Employment Contract":
            if "termination" not in content and "notice" not in content:
                issues.append({
                    "issue": "Missing termination clauses",
                    "severity": "Medium",
                    "suggestion": "Include termination and notice provisions",
                    "regulation": "ADGM Employment Regulations",
                    "section": "Employment Terms"
                })
        
        return issues

    def insert_comments(self, file_path: str, issues: List[Dict]) -> str:
        """Insert compliance comments into the document"""
        try:
            doc = docx.Document(file_path)
            
            # Add title for comments
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("ADGM COMPLIANCE REVIEW COMMENTS")
            title_run.bold = True
            title_run.font.size = docx.shared.Pt(14)
            
            # Add each issue as a comment
            for issue in issues:
                issue_para = doc.add_paragraph()
                issue_para.add_run(f"ISSUE: {issue['issue']}").bold = True
                issue_para.add_run(f"\nSeverity: {issue['severity']}")
                issue_para.add_run(f"\nSuggestion: {issue['suggestion']}")
                issue_para.add_run(f"\nRegulation: {issue['regulation']}")
                issue_para.add_run(f"\nSection: {issue['section']}")
                
                # Add spacing
                doc.add_paragraph()
            
            # Save as new file
            output_path = Path(file_path).with_suffix(".reviewed.docx")
            doc.save(str(output_path))
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error inserting comments: {str(e)}")
            raise

    def generate_comprehensive_report(self, documents: List[Dict], process_name: str) -> Dict:
        """Generate comprehensive compliance report"""
        try:
            # Get checklist verification
            uploaded_doc_names = [doc["filename"] for doc in documents]
            checklist_result = verify_checklist_completeness(uploaded_doc_names, process_name)
            
            # Count total issues
            total_issues = sum(len(doc.get("issues", [])) for doc in documents)
            high_severity_issues = sum(
                1 for doc in documents 
                for issue in doc.get("issues", []) 
                if issue.get("severity") == "High"
            )
            
            # Determine overall compliance status
            if total_issues == 0 and checklist_result["status"] == "complete":
                compliance_status = "Fully Compliant"
            elif high_severity_issues == 0:
                compliance_status = "Partially Compliant"
            else:
                compliance_status = "Non-Compliant"
            
            return {
                "process": process_name,
                "compliance_status": compliance_status,
                "checklist_verification": checklist_result,
                "documents_analyzed": len(documents),
                "total_issues_found": total_issues,
                "high_severity_issues": high_severity_issues,
                "document_details": documents,
                "recommendations": self._generate_recommendations(documents, checklist_result),
                "timestamp": "2025-08-10T18:30:00"
            }
            
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

    def _generate_recommendations(self, documents: List[Dict], checklist_result: Dict) -> List[str]:
        """Generate actionable recommendations"""
        recommendations = []
        
        # Missing documents
        if checklist_result["missing_documents"]:
            recommendations.append(
                f"Upload missing documents: {', '.join(checklist_result['missing_documents'])}"
            )
        
        # High severity issues
        high_issues = []
        for doc in documents:
            for issue in doc.get("issues", []):
                if issue.get("severity") == "High":
                    high_issues.append(f"{doc['filename']}: {issue['issue']}")
        
        if high_issues:
            recommendations.append("Address high severity issues immediately")
            recommendations.extend(high_issues[:3])  # Top 3 issues
        
        # General recommendations
        if checklist_result["completeness_percentage"] < 80:
            recommendations.append("Complete document package for better compliance")
        
        if not recommendations:
            recommendations.append("Documents appear compliant. Proceed with submission.")
        
        return recommendations 
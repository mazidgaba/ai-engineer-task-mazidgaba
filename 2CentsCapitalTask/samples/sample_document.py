#!/usr/bin/env python3
"""
Sample Document Generator for ADGM Corporate Agent Testing
This script creates a sample Articles of Association document for testing purposes.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

def create_sample_articles_of_association():
    """Create a sample Articles of Association document for testing"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('ARTICLES OF ASSOCIATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company name
    company_name = doc.add_heading('SAMPLE COMPANY LIMITED', 1)
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('These Articles of Association are made pursuant to the Companies Regulations of the United Arab Emirates.')
    
    # Add jurisdiction clause (INTENTIONALLY INCORRECT for testing)
    doc.add_paragraph()
    jurisdiction_heading = doc.add_heading('1. JURISDICTION', 2)
    jurisdiction_text = doc.add_paragraph()
    jurisdiction_text.add_run('This company shall be governed by the laws of the United Arab Emirates and shall be subject to the jurisdiction of the UAE Federal Courts.')
    
    # Add company objects
    doc.add_paragraph()
    objects_heading = doc.add_heading('2. COMPANY OBJECTS', 2)
    objects_text = doc.add_paragraph()
    objects_text.add_run('The objects for which the Company is established are:')
    
    objects_list = doc.add_paragraph()
    objects_list.add_run('a) To carry on any business as a general trading company\n')
    objects_list.add_run('b) To provide consulting services\n')
    objects_list.add_run('c) To invest in and manage other companies\n')
    objects_list.add_run('d) To do all such other things as are incidental or conducive to the attainment of the above objects')
    
    # Add share capital (missing governing law reference)
    doc.add_paragraph()
    capital_heading = doc.add_heading('3. SHARE CAPITAL', 2)
    capital_text = doc.add_paragraph()
    capital_text.add_run('The share capital of the Company shall be AED 1,000,000 divided into 1,000,000 ordinary shares of AED 1.00 each.')
    
    # Add directors
    doc.add_paragraph()
    directors_heading = doc.add_heading('4. DIRECTORS', 2)
    directors_text = doc.add_paragraph()
    directors_text.add_run('The Company shall have a minimum of one director and a maximum of ten directors.')
    
    # Add shareholders
    doc.add_paragraph()
    shareholders_heading = doc.add_heading('5. SHAREHOLDERS', 2)
    shareholders_text = doc.add_paragraph()
    shareholders_text.add_run('The Company shall have a minimum of one shareholder and a maximum of fifty shareholders.')
    
    # Add execution section (missing signature requirements)
    doc.add_paragraph()
    execution_heading = doc.add_heading('6. EXECUTION', 2)
    execution_text = doc.add_paragraph()
    execution_text.add_run('These Articles of Association shall come into effect upon registration with the relevant authorities.')
    
    # Add date
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.add_run('Dated: January 15, 2025')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save the document
    output_dir = "docs"
    os.makedirs(output_dir, exist_ok=True)
    
    filename = os.path.join(output_dir, "Sample_Articles_of_Association.docx")
    doc.save(filename)
    
    print(f"✅ Sample document created: {filename}")
    print("📝 This document contains intentional compliance issues for testing:")
    print("   - Incorrect jurisdiction (UAE instead of ADGM)")
    print("   - Missing governing law clause")
    print("   - Missing signature requirements")
    print("   - Missing ADGM-specific references")
    
    return filename

def create_sample_ubo_declaration():
    """Create a sample UBO Declaration document for testing"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('ULTIMATE BENEFICIAL OWNER DECLARATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company name
    company_name = doc.add_heading('SAMPLE COMPANY LIMITED', 1)
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('I, the undersigned, hereby declare the following information regarding the ultimate beneficial ownership of the above-named company.')
    
    # Add UBO details
    doc.add_paragraph()
    ubo_heading = doc.add_heading('1. ULTIMATE BENEFICIAL OWNER DETAILS', 2)
    
    # Add form fields
    doc.add_paragraph('Full Name: _________________________________')
    doc.add_paragraph('Date of Birth: _____________________________')
    doc.add_paragraph('Nationality: _______________________________')
    doc.add_paragraph('Passport Number: ___________________________')
    doc.add_paragraph('Address: ___________________________________')
    doc.add_paragraph('Percentage of Ownership: ___________________')
    
    # Add declaration
    doc.add_paragraph()
    declaration_heading = doc.add_heading('2. DECLARATION', 2)
    declaration_text = doc.add_paragraph()
    declaration_text.add_run('I hereby declare that the information provided above is true, accurate, and complete to the best of my knowledge.')
    
    # Add signature section (missing witness requirement)
    doc.add_paragraph()
    signature_heading = doc.add_heading('3. SIGNATURE', 2)
    doc.add_paragraph('Signature: _________________________________')
    doc.add_paragraph('Date: _____________________________________')
    
    # Save the document
    output_dir = "docs"
    os.makedirs(output_dir, exist_ok=True)
    
    filename = os.path.join(output_dir, "Sample_UBO_Declaration.docx")
    doc.save(filename)
    
    print(f"✅ Sample UBO document created: {filename}")
    print("📝 This document contains intentional compliance issues for testing:")
    print("   - Missing witness requirement")
    print("   - Missing ADGM jurisdiction reference")
    print("   - Incomplete signature section")
    
    return filename

if __name__ == "__main__":
    print("🏛️ Creating sample documents for ADGM Corporate Agent testing...")
    print()
    
    # Create sample documents
    articles_file = create_sample_articles_of_association()
    ubo_file = create_sample_ubo_declaration()
    
    print()
    print("🎯 Sample documents created successfully!")
    print("📁 Files saved in the 'docs' directory")
    print()
    print("🔍 You can now use these documents to test the ADGM Corporate Agent:")
    print("   1. Launch the application: python adgm_corporate_agent.py")
    print("   2. Upload the sample documents")
    print("   3. Select 'Company Incorporation' process")
    print("   4. Process documents to see compliance analysis")
    print()
    print("⚠️  Note: These documents contain intentional compliance issues")
    print("   to demonstrate the system's red flag detection capabilities.")
